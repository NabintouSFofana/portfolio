from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_line_break():
    return Document().add_paragraph()

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

# Header
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NABINTOU S. FOFANA")
run.bold = True
run.font.size = Pt(16)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("Plano, TX · 682-257-7208 · nabintousfofana@gmail.com").font.size = Pt(11)
contact.paragraph_format.space_after = Pt(2)

links = doc.add_paragraph()
links.alignment = WD_ALIGN_PARAGRAPH.CENTER
links.add_run("github.com/NabintouSFofana · linkedin.com/in/nabintousfofana").font.size = Pt(11)
links.paragraph_format.space_after = Pt(10)

# Add line
def add_hr():
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

add_hr()

# EDUCATION
edu_title = doc.add_paragraph()
run = edu_title.add_run("EDUCATION")
run.bold = True
run.font.size = Pt(12)
edu_title.paragraph_format.space_after = Pt(4)

edu = doc.add_paragraph("University of Texas at Dallas  |  B.S. Software Engineering  |  Expected May 2027")
edu.runs[0].font.size = Pt(11)
edu.paragraph_format.space_after = Pt(12)

# EXPERIENCE
exp_title = doc.add_paragraph()
run = exp_title.add_run("EXPERIENCE")
run.bold = True
run.font.size = Pt(12)
exp_title.paragraph_format.space_after = Pt(4)

role = doc.add_paragraph()
role_run = role.add_run("Freelance Developer — Schicgirl™")
role_run.bold = True
role_run.font.size = Pt(11)
role_run = role.add_run("  |  Mar 2024 – Present")
role_run.italic = True
role_run.font.size = Pt(11)
role.paragraph_format.space_after = Pt(4)

bullets = [
    "Built and manage a bilingual natural-hair platform (80+ pages, 38 HTML sources) from the ground up using hand-written HTML/CSS/JS, no framework. Deployed on GitHub Pages.",
    "Architected Supabase backend with row-level security across 11 tables to gate a paid membership, forum, and interactive studio. Single login opens everything at once.",
    "Caught and fixed production bugs: token expiry locking paying members out (refresh-before-expiry + retry), dropdown value mismatch causing silent data corruption, translated form fields storing different values per language."
]

for bullet_text in bullets:
    p = doc.add_paragraph(bullet_text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

doc.paragraphs[-1].paragraph_format.space_after = Pt(12)

# PROJECTS
proj_title = doc.add_paragraph()
run = proj_title.add_run("PROJECTS")
run.bold = True
run.font.size = Pt(12)
proj_title.paragraph_format.space_after = Pt(4)

# TerraScape
terra = doc.add_paragraph()
run = terra.add_run("TerraScape — Unity/C# Terrain Sandbox")
run.bold = True
run.font.size = Pt(11)
run = terra.add_run("  |  Computer Graphics  |  2026")
run.italic = True
run.font.size = Pt(11)
terra.paragraph_format.space_after = Pt(4)

terra_bullets = [
    "Built object placement system with Physics.Raycast for mouse picking. Implemented fractal L-system engine that generates procedural trees and space-filling curves (Koch, Sierpinski, dragon curve) from grammar rules.",
    "Owned team Git workflow for 4 engineers: set up repo structure, .gitignore, branch conventions, and onboarding docs so teammates with different Git experience could work without colliding on scene files."
]

for bullet_text in terra_bullets:
    p = doc.add_paragraph(bullet_text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

# InsightFlow
insight = doc.add_paragraph()
run = insight.add_run("InsightFlow — CSV Analytics Tool")
run.bold = True
run.font.size = Pt(11)
run = insight.add_run("  |  2026")
run.italic = True
run.font.size = Pt(11)
insight.paragraph_format.space_after = Pt(4)

insight_bullets = [
    "Flask app and CLI that turns messy CSVs into one-page PDF reports with stats, charts, and plain-English insights. Refactored a single 500-line script into reusable modules.",
    "Deployed to Render. Stack: Python, Flask, Pandas, Matplotlib, FPDF."
]

for bullet_text in insight_bullets:
    p = doc.add_paragraph(bullet_text, style='List Bullet')
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_after = Pt(5)

doc.paragraphs[-1].paragraph_format.space_after = Pt(4)

# Math Adventure
math = doc.add_paragraph()
run = math.add_run("Math Adventure — Spring Boot Platform")
run.bold = True
run.font.size = Pt(11)
run = math.add_run("  |  Software Engineering Course  |  2025")
run.italic = True
run.font.size = Pt(11)
math.paragraph_format.space_after = Pt(4)

math_bullet = "Came back solo after the initial team project to rebuild the backend: refactored utility classes into Spring services with dependency injection, converted responses to JSON DTOs with proper status codes, split 635-line frontend into 8 ES modules. Grew tests from 30 to 50 cases."
p = doc.add_paragraph(math_bullet, style='List Bullet')
p.runs[0].font.size = Pt(11)
p.paragraph_format.space_after = Pt(12)

# SKILLS
skills_title = doc.add_paragraph()
run = skills_title.add_run("SKILLS")
run.bold = True
run.font.size = Pt(12)
skills_title.paragraph_format.space_after = Pt(4)

skills_rows = [
    ("Languages", "Python, Java, JavaScript, C#, HTML, CSS, SQL"),
    ("Frameworks", "Flask, Spring Boot, Supabase, React (learning)"),
    ("Tools", "Git, GitHub, Unity, Google Apps Script, Render, Pandas, Matplotlib, FPDF, Linux")
]

for label, text in skills_rows:
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(f"  {text}")
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

doc.save("Nabintou_Fofana_resume.docx")
print("DOCX created")
